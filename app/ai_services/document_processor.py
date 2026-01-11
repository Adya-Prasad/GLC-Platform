"""
Document Processor for Sustainability-Linked Loan Analysis
Extracts text from PDF/DOCX sustainability reports and annual reports.
"""

import logging
import re
from pathlib import Path
from typing import Optional, List
from dataclasses import dataclass

from app.ai_services.config import settings

logger = logging.getLogger(__name__)


@dataclass
class DocumentChunk:
    """A chunk of document text with metadata."""
    text: str
    chunk_index: int
    page_number: Optional[int] = None
    source_file: str = ""
    doc_type: str = ""


@dataclass
class ProcessedDocument:
    """Result of document processing."""
    filename: str
    doc_type: str
    full_text: str
    chunks: List[DocumentChunk]
    page_count: int
    word_count: int
    extraction_method: str


class DocumentProcessor:
    """
    Processes sustainability reports and annual reports for RAG analysis.
    Supports PDF, DOCX, and TXT formats.
    """
    
    CHUNK_SIZE = settings.TEXT_CHUNK_SIZE
    CHUNK_OVERLAP = settings.TEXT_CHUNK_OVERLAP
    
    def __init__(self):
        self.logger = logging.getLogger(f"{__name__}.DocumentProcessor")
    
    def process_document(
        self, 
        filepath: str, 
        doc_type: str = "sustainability_report"
    ) -> Optional[ProcessedDocument]:
        """
        Process a document and return structured result.
        """
        path = Path(filepath)
        
        if not path.exists():
            self.logger.error(f"File not found: {filepath}")
            return None
        
        ext = path.suffix.lower()
        
        # Extract text based on file type
        if ext == '.pdf':
            text, page_count, method = self._extract_pdf(filepath)
        elif ext in ['.docx', '.doc']:
            text, page_count, method = self._extract_docx(filepath)
        elif ext == '.txt':
            text, page_count, method = self._extract_txt(filepath)
        else:
            self.logger.warning(f"Unsupported file type: {ext}")
            return None
        
        if not text or len(text.strip()) < 50:
            self.logger.warning(f"No meaningful text extracted from {filepath}")
            return None
        
        # Clean and normalize text
        text = self._clean_text(text)
        
        # Create chunks
        chunks = self._create_chunks(text, path.name, doc_type)
        
        return ProcessedDocument(
            filename=path.name,
            doc_type=doc_type,
            full_text=text,
            chunks=chunks,
            page_count=page_count,
            word_count=len(text.split()),
            extraction_method=method
        )
    
    def _extract_pdf(self, filepath: str) -> tuple:
        """Extract text from PDF."""
        text = ""
        page_count = 0
        method = "pdfminer"
        
        try:
            from pdfminer.high_level import extract_text as pdfminer_extract
            from pdfminer.pdfpage import PDFPage
            
            text = pdfminer_extract(filepath)
            
            with open(filepath, 'rb') as f:
                page_count = len(list(PDFPage.get_pages(f)))
            
            if text and len(text.strip()) > 100:
                return text.strip(), page_count, method
        except Exception as e:
            self.logger.warning(f"pdfminer failed: {e}")
        
        # Fallback to PyPDF2
        try:
            import PyPDF2
            method = "PyPDF2"
            
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                page_count = len(reader.pages)
                pages = [page.extract_text() or "" for page in reader.pages]
                text = "\n\n".join(pages)
            
            if text and len(text.strip()) > 100:
                return text.strip(), page_count, method
        except Exception as e:
            self.logger.warning(f"PyPDF2 failed: {e}")
        
        return text, page_count, method
    
    def _extract_docx(self, filepath: str) -> tuple:
        """Extract text from DOCX file."""
        try:
            from docx import Document
            
            doc = Document(filepath)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n\n".join(paragraphs)
            page_count = max(1, len(text.split()) // 500)
            
            return text, page_count, "python-docx"
        except Exception as e:
            self.logger.error(f"DOCX extraction failed: {e}")
            return "", 0, "failed"
    
    def _extract_txt(self, filepath: str) -> tuple:
        """Extract text from plain text file."""
        try:
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            page_count = max(1, len(text.split()) // 500)
            return text, page_count, "plaintext"
        except Exception as e:
            self.logger.error(f"TXT extraction failed: {e}")
            return "", 0, "failed"
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove page numbers
        text = re.sub(r'\n\s*\d+\s*\n', '\n', text)
        # Normalize line breaks
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()
    
    def _create_chunks(
        self, 
        text: str, 
        filename: str, 
        doc_type: str
    ) -> List[DocumentChunk]:
        """Split text into overlapping chunks."""
        words = text.split()
        chunks = []
        
        start_idx = 0
        chunk_index = 0
        
        while start_idx < len(words):
            end_idx = min(start_idx + self.CHUNK_SIZE, len(words))
            chunk_words = words[start_idx:end_idx]
            chunk_text = ' '.join(chunk_words)
            
            chunks.append(DocumentChunk(
                text=chunk_text,
                chunk_index=chunk_index,
                source_file=filename,
                doc_type=doc_type
            ))
            
            # Move forward with overlap
            start_idx = end_idx - self.CHUNK_OVERLAP if end_idx < len(words) else len(words)
            chunk_index += 1
        
        self.logger.info(f"Created {len(chunks)} chunks from {filename}")
        return chunks


# Singleton instance
document_processor = DocumentProcessor()
